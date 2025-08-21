import os
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
 
def excel_to_word_converter(excel_path, output_path):
    """
    Convert Excel file to Word document with proper formatting
    """
    try:
        # Load the workbook
        wb = load_workbook(excel_path, data_only=True)
       
        # Create a new Word document
        doc = Document()
       
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)
       
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
           
            # Add a heading for the sheet
            if len(wb.sheetnames) > 1:  # Only add heading if multiple sheets
                heading = doc.add_heading(sheet_name, level=1)
                heading.style.font.name = 'Arial'
                heading.style.font.bold = True
           
            # Create a table in Word
            table = doc.add_table(rows=1, cols=len(ws[1]))
            table.style = 'Table Grid'
           
            # Add headers (first row)
            hdr_cells = table.rows[0].cells
            for i, cell in enumerate(ws[1], 1):
                hdr_cells[i-1].text = str(cell.value) if cell.value is not None else ""
                # Format header
                paragraph = hdr_cells[i-1].paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hdr_cells[i-1].paragraphs[0].paragraph_format.space_after = Pt(6)
           
            # Add data rows
            for row_idx, row in enumerate(ws.iter_rows(min_row=2), 2):
                row_cells = table.add_row().cells
                for col_idx, cell in enumerate(row, 1):
                    if cell.value is not None:
                        # Format numbers with 2 decimal places if they're numbers
                        if isinstance(cell.value, (int, float)):
                            # Check if the cell has a number format that indicates currency
                            if ws.cell(row=row_idx, column=col_idx).number_format and \
                               any(x in str(ws.cell(row=row_idx, column=col_idx).number_format) for x in ['$', '€', '£', '¥']):
                                row_cells[col_idx-1].text = f"${cell.value:,.2f}"
                            else:
                                row_cells[col_idx-1].text = f"{cell.value:,.2f}"
                           
                            # Right-align numeric cells
                            paragraph = row_cells[col_idx-1].paragraphs[0]
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            row_cells[col_idx-1].text = str(cell.value)
                           
                        # Apply cell formatting
                        paragraph = row_cells[col_idx-1].paragraphs[0]
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        run.font.name = 'Arial'
                        run.font.size = Pt(9)
                       
                        # Apply cell background color if exists
                        if cell.fill and cell.fill.start_color and cell.fill.start_color.rgb:
                            # Convert RGB to hex and set shading
                            rgb = cell.fill.start_color.rgb[2:]  # Remove alpha channel
                            if rgb != '000000':  # Skip black
                                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{rgb}"')
                                row_cells[col_idx-1]._tc.get_or_add_tcPr().append(shading_elm)
           
            # Auto-fit table columns
            table.autofit = True
           
            # Add space after each table
            doc.add_paragraph()
       
        # Save the document
        doc.save(output_path)
        print(f"✅ Excel to Word conversion successful: {output_path}")
       
    except Exception as e:
        print(f"❌ Error converting Excel to Word: {e}")
        raise